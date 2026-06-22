"""Turn a generated article (markdown) into reading stats and downloadable files.

The heavy converters (PDF via PyMuPDF, DOCX via python-docx) are imported lazily
so the rest of the app still boots if those wheels are missing in some
environment; the server surfaces a clean 503 instead of failing at import time.
"""

import io
import re

# Average adult reading speed for prose; used for the "x min read" badge.
WORDS_PER_MINUTE = 200


def word_count(content: str) -> int:
    """Count words in the article, ignoring markdown punctuation noise."""
    return len(re.findall(r"\b\w[\w'-]*\b", content))


def reading_time_minutes(content: str) -> int:
    """Estimated minutes to read the article, rounded up to at least 1."""
    words = word_count(content)
    if words == 0:
        return 0
    return max(1, round(words / WORDS_PER_MINUTE))


def reading_stats(content: str) -> dict:
    """Word count + reading time, ready to drop into a JSON response."""
    return {
        "words": word_count(content),
        "reading_time_minutes": reading_time_minutes(content),
    }


# ---- File exports ----

class ExportUnavailable(RuntimeError):
    """Raised when the library needed for a given export format isn't installed."""


def _strip_markdown(content: str) -> str:
    """Best-effort plain-text rendering of markdown for the simple exporters."""
    text = re.sub(r"`{1,3}([^`]*)`{1,3}", r"\1", content)          # inline/code fences
    text = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", text)                # images
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)            # links -> text
    text = re.sub(r"^\s{0,3}#{1,6}\s*", "", text, flags=re.MULTILINE)  # heading hashes
    text = re.sub(r"(\*\*|__|\*|_)", "", text)                      # bold/italic markers
    text = re.sub(r"^\s{0,3}>\s?", "", text, flags=re.MULTILINE)    # blockquotes
    return text


def _wrap_line(text: str, font, fontsize: float, max_width: float) -> list[str]:
    """Greedy word-wrap a single line to fit max_width at the given font size."""
    words = text.split()
    if not words:
        return [""]
    lines, current = [], words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if font.text_length(candidate, fontsize) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def to_pdf(content: str, title: str) -> bytes:
    """Render the article to a simple PDF and return the bytes.

    Draws text line by line with explicit positioning rather than relying on
    insert_textbox's overflow/reflow behaviour, which could silently leave a
    page blank in some environments.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ExportUnavailable("PDF export requires PyMuPDF (fitz).") from e

    fontname, fontsize, leading = "helv", 11, 15
    margin = 54
    font = fitz.Font(fontname)
    # The base-14 fonts only cover Latin-1; map anything else to a close ASCII
    # equivalent so smart quotes / dashes / bullets from the LLM still render.
    body = _strip_markdown(content)
    text = f"{title}\n\n{body}" if title else body
    text = (
        text.replace("‘", "'").replace("’", "'")
        .replace("“", '"').replace("”", '"')
        .replace("–", "-").replace("—", "-")
        .replace("•", "-").replace("…", "...")
    )
    text = text.encode("latin-1", "replace").decode("latin-1")

    doc = fitz.open()
    page = doc.new_page()
    width, height = page.rect.width, page.rect.height
    max_width = width - 2 * margin
    y = margin

    for raw_line in text.split("\n"):
        for line in _wrap_line(raw_line, font, fontsize, max_width):
            if y + leading > height - margin:
                page = doc.new_page()
                y = margin
            page.insert_text((margin, y), line, fontsize=fontsize, fontname=fontname)
            y += leading

    out = doc.tobytes()
    doc.close()
    return out


def to_docx(content: str, title: str) -> bytes:
    """Render the article to a .docx and return the bytes."""
    try:
        from docx import Document
    except ImportError as e:
        raise ExportUnavailable("DOCX export requires python-docx.") from e

    document = Document()
    if title:
        document.add_heading(title, level=0)

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(_strip_markdown(heading.group(2)), level=level)
        else:
            document.add_paragraph(_strip_markdown(stripped))

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
